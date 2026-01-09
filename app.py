import streamlit as st
import json
import datetime
from openai import OpenAI
from serpapi import GoogleSearch
import io

# 👇 地图相关库
try:
    import folium
    from streamlit_folium import st_folium
except ImportError:
    st.error("⚠️ 缺少地图库！请运行: pip install folium streamlit-folium")
    st.stop()

# 👇 用于生成 Word 文档
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn 
except ImportError:
    st.error("请先安装 python-docx 库: pip install python-docx")
    st.stop()

# ==========================================
# 1. 页面基础设置
# ==========================================
st.set_page_config(page_title="AI 智能旅行管家 (自主思考版)", page_icon="🧠", layout="wide")

# 👇 你的 API KEY (保持不变)
DEEPSEEK_API_KEY = "YOUR_DEEPSEEK_API_KEY_HERE" 
SERPAPI_API_KEY = "YOUR_SERPAPI_API_KEY_HERE"

# 初始化客户端
try:
    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
except Exception as e:
    st.error("API Key 配置有误，请检查代码。")
    st.stop()

# ==========================================
# 2. 定义系统核心指令
# ==========================================
today = datetime.date.today().strftime("%Y-%m-%d")

SYSTEM_PROMPT = f"""
You are an **Autonomous AI Travel Agent**.
📅 **TODAY'S DATE**: {today}
⚠️ **CRITICAL**: All date calculations (e.g., "next month", "next year") MUST be based on {today}.
📍 **DEFAULT ORIGIN**: Kuala Lumpur (KUL) (Unless user specifies otherwise).

【🔴 CORE IDENTITY: VISUAL PLANNER】
You must plan the route logically.
**AFTER** creating the text itinerary, you **MUST** call the tool `generate_map_with_traffic` to visualize the route.
Pass the list of locations in order (e.g., ["Hotel", "Attraction A", "Restaurant", "Attraction B"]).

【🔴 CORE IDENTITY: CHAIN OF THOUGHT (CoT)】
Before answering or calling tools, you must **THINK** in steps.
1. **Analyze**: What is the user's *real* goal?
2. **Plan**: What information is missing? What tools do I need?
3. **Execute**: Call tools.
4. **Verify & Self-Correct (CRITICAL)**:
   - If `search_flights` returns "No flights", **DO NOT** give up. 
   - **THINK**: "Is the date too far ahead?"
   - **ACTION**: Use `search_general_web`.

【🔴 RULE 0: MAP GENERATION POLICY (SPEED OPTIMIZATION)】
- **DEFAULT BEHAVIOR**: PROHIBITED to generate maps automatically.
- **EXCEPTION**: ONLY call `generate_map_with_traffic` if the user **EXPLICITLY** asks for it (e.g., "show map", "visualize route", "画个地图", "怎么走").
- **Reasoning**: Generating maps is slow. Prioritize quick text responses first.

【🔴 RULE 1: DYNAMIC LANGUAGE SWITCHING】
- User speaks **Chinese** -> Reply in **Chinese**.
- User speaks **English** -> Reply in **English**.
- User speaks **Malay/Rojak** -> Reply in **Malay/Manglish**.

【🔴 RULE 2: FORMATTING (Clean Markdown)】
- **DO NOT use HTML.** Use standard **Markdown**.
- Use **Bold** for emphasis.
- Use **Lists** for itinerary steps.
"""

# ==========================================
# 🟢 Part 4: 工具函数 (保持完整)
# ==========================================

# 1. 机票搜索
def search_flights(origin, destination, date, return_date):
    st.toast(f"✈️ Checking Flights: {origin}->{destination} ({date})") 
    print(f"[后台] 查机票 {origin}-{destination} ({date})")
    params = {
        "engine": "google_flights", "departure_id": origin, "arrival_id": destination, 
        "outbound_date": date, "return_date": return_date, "currency": "MYR", "hl": "en", 
        "api_key": SERPAPI_API_KEY, "type": "1"
    }
    try:
        res = GoogleSearch(params).get_dict()
        if "best_flights" not in res: 
            return f"RESULT: No specific flights found for {date}. (HINT: Date might be too far ahead?)"
        f = res['best_flights'][0]
        return json.dumps({
            "date": date, "airline": f['flights'][0]['airline'], 
            "price_per_adult": f['price'], "duration": f['total_duration']
        })
    except: return f"Error searching flights for {date}"

# 2. 酒店搜索
def search_hotels(city, check_in_date, check_out_date, adults):
    st.toast(f"🏨 Checking Hotels: {city}")
    search_type = "Vacation Rentals" if adults > 2 else "Hotels"
    params = {"engine": "google_hotels", "q": f"{city} {search_type}", "check_in_date": check_in_date, "check_out_date": check_out_date, "adults": adults, "currency": "MYR", "hl": "en", "gl": "my", "api_key": SERPAPI_API_KEY}
    try:
        res = GoogleSearch(params).get_dict()
        hotels = []
        if "properties" in res:
            for h in res["properties"][:3]:
                name = h.get("name")
                price = h.get("rate_per_night", {}).get("lowest", "N/A")
                hotels.append(f"- {name} | Price: {price}")
        return "\n".join(hotels) if hotels else "No hotels found"
    except: return "Error searching hotels"

# 3. 景点搜索
def search_attractions(city, keyword=None):
    st.toast(f"🎡 Checking Sights: {city}")
    q = f"top sights in {city}" if not keyword else f"best {keyword} in {city}"
    params = {"engine": "google_maps", "q": q, "type": "search", "hl": "en", "api_key": SERPAPI_API_KEY}
    try:
        res = GoogleSearch(params).get_dict()
        results = []
        for r in res.get("local_results", [])[:5]:
            title = r.get('title', 'Unknown')
            rating = r.get('rating', 'N/A')
            results.append(f"- {title} ({rating}⭐)")
        return "\n".join(results)
    except: return "Error searching attractions"

# 4. 美食搜索
def search_restaurants(city, food_type="local food"):
    st.toast(f"🍜 Checking Food: {city}")
    params = {"engine": "google_maps", "q": f"best {food_type} in {city}", "type": "search", "hl": "en", "api_key": SERPAPI_API_KEY}
    try:
        res = GoogleSearch(params).get_dict()
        return "\n".join([f"- {r.get('title')} ({r.get('rating')}⭐)" for r in res.get("local_results", [])[:5]])
    except: return "Error searching food"

# 5. 通用搜索
def search_general_web(query):
    st.toast(f"🧠 Brain: Googling '{query}'...")
    params = {"engine": "google", "q": query, "hl": "en", "gl": "my", "api_key": SERPAPI_API_KEY}
    try:
        res = GoogleSearch(params).get_dict()
        snippets = [f"- {r.get('title')}: {r.get('snippet')}" for r in res.get("organic_results", [])[:3]]
        return "\n".join(snippets) if snippets else "No web results found."
    except: return "Web search error."

# --- 辅助函数：获取经纬度 ---
def get_coordinates(location):
    params = {"engine": "google_maps", "q": location, "type": "search", "api_key": SERPAPI_API_KEY}
    try:
        res = GoogleSearch(params).get_dict()
        if "local_results" in res and res["local_results"]:
            gps = res["local_results"][0].get("gps_coordinates", {})
            return gps.get("latitude"), gps.get("longitude"), res["local_results"][0].get("title", location)
    except: pass
    return None, None, location

# 2. 核心升级：分组显示 (驾车独立，公交+走路合并)
def get_directions(start_lat, start_lng, end_lat, end_lng):
    start = f"{start_lat},{start_lng}"
    end = f"{end_lat},{end_lng}"
    
    results = {
        "0": {"icon": "🚗", "text": "驾车", "time": "N/A", "details": ""},
        "3": {"icon": "🚇", "text": "公交", "time": "N/A", "details": ""},
        "2": {"icon": "🚶", "text": "步行", "time": "N/A", "details": ""}
    }
    
    for mode_code, info in results.items():
        params = {
            "engine": "google_maps_directions",
            "start_coords": start,
            "end_coords": end,
            "travel_mode": mode_code, 
            "api_key": SERPAPI_API_KEY
        }
        try:
            res = GoogleSearch(params).get_dict()
            if "directions" in res and res["directions"]:
                route = res["directions"][0]
                duration = route.get("formatted_duration", "N/A")
                results[mode_code]["time"] = duration
                
                if mode_code == "3" and "legs" in route:
                    steps = route["legs"][0].get("steps", [])
                    transit_segments = []
                    for step in steps:
                        if step.get("travel_mode") == "TRANSIT" and "transit_details" in step:
                            td = step["transit_details"]
                            line_name = td.get("line", {}).get("short_name") or td.get("line", {}).get("name") or "Bus"
                            transit_segments.append(f"**{line_name}**") 
                    
                    if transit_segments:
                        results[mode_code]["details"] = f" ➤ [{' > '.join(transit_segments)}]"
        except: 
            pass 
            
    line1 = f"🚗 **驾车**: {results['0']['time']}" if results["0"]["time"] != "N/A" else "🚗 驾车: 无法到达"
    transit_str = f"🚇 **公交**: {results['3']['time']}{results['3']['details']}" if results["3"]["time"] != "N/A" else "🚇 公交: N/A"
    walk_str = f"🚶 **步行**: {results['2']['time']}" if results["2"]["time"] != "N/A" else "🚶 步行: N/A"
    
    line2 = f"{transit_str} &nbsp;&nbsp;|&nbsp;&nbsp; {walk_str}"
    return f"{line1}\n\n{line2}"

# 3. 地图生成
def generate_map_with_traffic(locations_list):
    if len(locations_list) < 1: return "Need at least 1 location."
    
    st.toast(f"🗺️ Calculating optimized routes for: {', '.join(locations_list)}...")
    
    coords = []
    for loc in locations_list:
        lat, lng, real_name = get_coordinates(loc)
        if lat and lng:
            coords.append((lat, lng, real_name))
            
    if not coords: return "Could not find coordinates."

    m = folium.Map(location=[coords[0][0], coords[0][1]], zoom_start=13)
    
    traffic_info = []
    
    for i in range(len(coords)):
        lat, lng, name = coords[i]
        
        folium.Marker(
            [lat, lng], 
            popup=name, 
            tooltip=f"{i+1}. {name}",
            icon=folium.Icon(color='red' if i==0 else 'blue', icon='info-sign')
        ).add_to(m)
        
        if i < len(coords) - 1:
            next_lat, next_lng, next_name = coords[i+1]
            travel_str = get_directions(lat, lng, next_lat, next_lng)
            traffic_info.append(f"🚩 **{name} ➡️ {next_name}**")
            traffic_info.append(travel_str) 
            folium.PolyLine(locations=[[lat, lng], [next_lat, next_lng]], color="blue", weight=4, opacity=0.6, dash_array='10').add_to(m)

    st.session_state["map_data"] = m
    st.session_state["traffic_data"] = "\n\n".join(traffic_info)
    
    return "Map Generated!"

# 👇 专门处理删除逻辑的函数 (修复 Delete Button 的关键！)
def delete_chat_history(index):
    if 0 <= index < len(st.session_state["chat_history"]):
        st.session_state["chat_history"].pop(index)
        
        if st.session_state["current_chat_id"] == index:
            st.session_state["current_chat_id"] = None
            st.session_state["messages"] = [
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "assistant", "content": "Chat deleted."}
            ]
            st.session_state["download_buffer"] = None
            st.session_state["map_data"] = None 
            st.session_state["traffic_data"] = None
        elif st.session_state["current_chat_id"] is not None and st.session_state["current_chat_id"] > index:
            st.session_state["current_chat_id"] -= 1

# 👇 生成 Word 文档函数
def create_word_doc(content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_heading('Travel Itinerary (AI Generated)', 0)
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('**') and line.endswith('**'): doc.add_paragraph().add_run(line.replace('**', '')).bold = True
        elif line.startswith('- '): doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        else: doc.add_paragraph(line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 👇 保存函数
def save_itinerary(content):
    doc_buffer = create_word_doc(content)
    st.session_state["download_buffer"] = doc_buffer
    if st.session_state["current_chat_id"] is not None:
        chat_id = st.session_state["current_chat_id"]
        if 0 <= chat_id < len(st.session_state["chat_history"]):
            st.session_state["chat_history"][chat_id]["itinerary_content"] = content
    try:
        with open("My_Trip_Plan.txt", "w", encoding="utf-8") as f: f.write(content)
    except: pass
    return "✅ Itinerary saved! Check the sidebar to download."

# 工具列表 (Tools List)
tools = [
    {"type": "function", "function": {"name": "search_flights", "description": "Search flights", "parameters": {"type": "object", "properties": {"origin": {"type": "string"}, "destination": {"type": "string"}, "date": {"type": "string"}, "return_date": {"type": "string"}}, "required": ["origin", "destination", "date", "return_date"]}}},
    {"type": "function", "function": {"name": "search_hotels", "description": "Search hotels", "parameters": {"type": "object", "properties": {"city": {"type": "string"}, "check_in_date": {"type": "string"}, "check_out_date": {"type": "string"}, "adults": {"type": "integer"}}, "required": ["city", "check_in_date", "check_out_date", "adults"]}}},
    {"type": "function", "function": {"name": "search_attractions", "description": "Search attractions", "parameters": {"type": "object", "properties": {"city": {"type": "string"}, "keyword": {"type": "string"}}, "required": ["city"]}}},
    {"type": "function", "function": {"name": "search_restaurants", "description": "Search restaurants", "parameters": {"type": "object", "properties": {"city": {"type": "string"}, "food_type": {"type": "string"}}, "required": ["city"]}}},
    {"type": "function", "function": {"name": "search_general_web", "description": "Search Google for general info", "parameters": {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]}}},
    {"type": "function", "function": {"name": "save_itinerary", "description": "Generate Word document", "parameters": {"type": "object", "properties": {"content": {"type": "string"}}, "required": ["content"]}}},
    {"type": "function", "function": {"name": "generate_map_with_traffic", "description": "Generate a map. ⚠️ ONLY use this if user explicitly asks for 'map' or 'route visualization'. Do not use by default.", "parameters": {"type": "object", "properties": {"locations_list": {"type": "array", "items": {"type": "string"}}}, "required": ["locations_list"]}}}
]

# ==========================================
# 3. 状态管理 (Session State)
# ==========================================

if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = []
if "current_chat_id" not in st.session_state:
    st.session_state["current_chat_id"] = None
if "messages" not in st.session_state:
    st.session_state["messages"] = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "assistant", "content": "Hello! I am your Autonomous AI Agent. Where are we going today?"}
    ]
if "download_buffer" not in st.session_state: 
    st.session_state["download_buffer"] = None
if "map_data" not in st.session_state: 
    st.session_state["map_data"] = None
if "traffic_data" not in st.session_state: 
    st.session_state["traffic_data"] = None

# ==========================================
# 🔴 Sidebar UI (侧边栏 - 修复版)
# ==========================================
with st.sidebar:
    st.title("🗂️ Control Panel")
    
    if st.button("➕ New Chat", use_container_width=True, type="primary"):
        # 自动保存旧对话
        if len(st.session_state["messages"]) > 2:
            user_first_msg = "New Chat"
            for msg in st.session_state["messages"]:
                if msg["role"] == "user":
                    content = msg.get("content", "")
                    user_first_msg = content[:15] + "..."
                    break
            
            if st.session_state["current_chat_id"] is None:
                st.session_state["chat_history"].insert(0, {
                    "title": user_first_msg,
                    "messages": st.session_state["messages"],
                    "itinerary_content": None
                })
        
        # 重置
        st.session_state["messages"] = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "assistant", "content": "Hello! Where are we going today?"}
        ]
        st.session_state["current_chat_id"] = None
        st.session_state["download_buffer"] = None
        st.session_state["map_data"] = None 
        st.session_state["traffic_data"] = None 
        st.rerun()

    # 下载按钮
    if st.session_state.get("download_buffer"):
        st.markdown("---")
        st.success("✅ 行程单已生成！")
        st.download_button(
            label="📥 下载 Word 行程单 (.docx)",
            data=st.session_state["download_buffer"],
            file_name="My_Trip_Plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )

    st.markdown("---")
    st.subheader("🕒 History")

    if not st.session_state["chat_history"]:
        st.caption("No history yet")
    
    # 历史记录列表 (修复版 - 使用 on_click)
    for i, chat in enumerate(st.session_state["chat_history"]):
        col1, col2 = st.columns([0.85, 0.15]) 
        with col1:
            if st.button(f"💬 {chat['title']}", key=f"hist_{i}", use_container_width=True):
                if len(st.session_state["messages"]) > 2 and st.session_state["current_chat_id"] is None:
                     temp_title = "Unsaved"
                     for m in st.session_state["messages"]:
                         if m["role"] == "user":
                             temp_title = m.get("content", "")[:15] + "..."
                             break
                     st.session_state["chat_history"].insert(0, {
                         "title": temp_title, 
                         "messages": st.session_state["messages"],
                         "itinerary_content": None
                     })

                st.session_state["messages"] = chat["messages"]
                st.session_state["current_chat_id"] = i
                
                # 恢复内容
                saved_content = chat.get("itinerary_content")
                if saved_content:
                    st.session_state["download_buffer"] = create_word_doc(saved_content)
                else:
                    st.session_state["download_buffer"] = None
                
                st.rerun()
        
        with col2:
            # 🔴 使用 callback 解决删除按钮问题
            st.button("🗑️", key=f"del_{i}", on_click=delete_chat_history, args=(i,))

# ==========================================
# 5. 聊天主界面 (修复：即时显示历史记录)
# ==========================================
st.title("🧠 My AI Travel Agent (Map Edition)")

# 1. 聊天记录显示
for msg in st.session_state["messages"]:
    if isinstance(msg, dict):
        role = msg["role"]
        content = msg.get("content")
    else:
        role = msg.role
        content = msg.content
    
    if role == "system": continue

    if content:
        if role == "user": st.chat_message("user").write(content)
        elif role == "assistant": st.chat_message("assistant").write(content)

# [UI] 地图显示区域
if st.session_state.get("map_data"):
    with st.container():
        st.markdown("### 🗺️ 路线地图 & 交通时间")
        if st.session_state.get("traffic_data"):
            with st.expander("🚗 查看详细交通耗时 (驾车 vs 公交)", expanded=True):
                st.markdown(st.session_state["traffic_data"])
        try:
            from streamlit_folium import st_folium
            st_folium(st.session_state["map_data"], width=700, height=400, returned_objects=[])
        except ImportError:
            st.error("⚠️ 缺少地图组件")

# ==========================================
# 🔴 核心修改区域：把 "接收输入" 和 "AI回答" 分开
# ==========================================

# 2. 接收用户输入
if prompt := st.chat_input("Say something..."):
    # 存入消息
    st.session_state["messages"].append({"role": "user", "content": prompt})
    
    # 🟢 自动建档逻辑
    if st.session_state["current_chat_id"] is None:
        st.session_state["chat_history"].insert(0, {
            "title": prompt[:15] + "...", 
            "messages": st.session_state["messages"],
            "itinerary_content": None
        })
        st.session_state["current_chat_id"] = 0
    else:
        st.session_state["chat_history"][st.session_state["current_chat_id"]]["messages"] = st.session_state["messages"]
    
    # 🔥 关键修改：输入完立刻刷新！
    # 这样侧边栏会马上更新，显示出刚才新建的 Chat
    st.rerun()

# 3. AI 思考循环 (逻辑：如果最后一条消息是 user，说明 AI 还没回，需要干活)
# 注意：这里不再缩进在 if prompt 里面了，而是独立出来的
if st.session_state["messages"] and st.session_state["messages"][-1]["role"] == "user":
    
    with st.chat_message("assistant"):
        status_container = st.status("🧠 Agent is thinking...", expanded=True)
        messages = st.session_state["messages"]
        
        while True:
            response = client.chat.completions.create(model="deepseek-chat", messages=messages, tools=tools)
            msg = response.choices[0].message
            
            if msg.tool_calls:
                messages.append(msg)
                
                should_rerun = False 
                
                for tool in msg.tool_calls:
                    fn, args = tool.function.name, json.loads(tool.function.arguments)
                    status_container.write(f"👉 Action: **{fn}**")
                    
                    res = None
                    if fn == "search_flights": res = search_flights(args["origin"], args["destination"], args["date"], args.get("return_date"))
                    elif fn == "search_hotels": res = search_hotels(args["city"], args["check_in_date"], args.get("check_out_date"), args.get("adults", 1))
                    elif fn == "search_attractions": res = search_attractions(args["city"], args.get("keyword"))
                    elif fn == "search_restaurants": res = search_restaurants(args["city"], args.get("food_type"))
                    elif fn == "search_general_web": res = search_general_web(args["query"])
                    elif fn == "save_itinerary": res = save_itinerary(args["content"]); status_container.write("💾 Saved!")
                    
                    elif fn == "generate_map_with_traffic": 
                        res = generate_map_with_traffic(args["locations_list"])
                        status_container.write("🗺️ Map Drawn!")
                        should_rerun = True 
                    
                    messages.append({"role": "tool", "tool_call_id": tool.id, "content": str(res)})
                
                if should_rerun:
                    st.rerun()

            else:
                final_content = msg.content
                status_container.update(label="✅ Response Ready", state="complete", expanded=False)
                st.markdown(final_content)
                st.session_state["messages"].append({"role": "assistant", "content": final_content})
                
                # 同步更新到历史记录
                if st.session_state["current_chat_id"] is not None:
                     st.session_state["chat_history"][st.session_state["current_chat_id"]]["messages"] = st.session_state["messages"]
                break