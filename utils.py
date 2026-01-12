import io
import streamlit as st

# 👇 用于生成 Word 文档 (保持原本的 try-except 逻辑)
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn 
except ImportError:
    st.error("请先安装 python-docx 库: pip install python-docx")
    st.stop()

def create_word_doc(content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_heading('Travel Itinerary (AI Generated)', 0)
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('**') and line.endswith('**'): doc.add_paragraph().add_run(line.replace('**', '')).bold = True
        elif line.startswith('- '): doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        else: doc.add_paragraph(line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer